from datetime import datetime
import docx

class docManager:
    def __init__(self) -> None:
        self.document   = docx.Document("documents\\model.docx")
        self.docName    = ''
    
    def replaceKeys(self, keyPack):
        for keyGroup in keyPack:
            for key in keyGroup:
                for paragraph in self.document.paragraphs:
                        docTag = '%' + keyGroup[key]["tag"] + '%'
                        if (docTag) in paragraph.text:
                            paragraph = docManager.overwriteParagraphLines(paragraph, docTag, keyGroup[key]["value"])
                        if '%today%' in paragraph.text:
                            paragraph = docManager.overwriteParagraphLines(paragraph, '%today%', docManager.writeDate(datetime.now()))
                        if '%timestamp%' in paragraph.text:
                            paragraph = docManager.overwriteParagraphLines(paragraph, '%timestamp%', docManager.writeTime(datetime.now()))
                        if '%writedate%' in paragraph.text:
                            paragraph = docManager.overwriteParagraphLines(paragraph, '%writedate%', docManager.writeFormattedDate(datetime.now()))
                            
    def overwriteParagraphLines(paragraph, tag, value):
        inline = paragraph.runs
        for i in range(len(inline)):
            if tag in inline[i].text:
                inline[i].text = inline[i].text.replace(tag, value)
        return paragraph
    
    def setDocumentName(self, docName):
        self.docName = docName
        
    def saveDocument(self):
        self.document.save("documents\\{0}.docx".format(self.docName))
    
    def writeTime(date):
        formatHour      = str(date.hour) if date.hour > 9 else '0%d' % date.hour
        formatMinute    = str(date.minute) if date.minute > 9 else '0%d' % date.minute
        
        return '%s:%s' % (formatHour, formatMinute)
    
    def writeDate(date):
        formatDay   = str(date.day) if date.day > 9 else '0%d' % date.day
        formatMonth = str(date.month) if date.month > 9 else '0%d' % date.month
        
        return '%s/%s/%d' % (formatDay, formatMonth, date.year)
    
    def writeFormattedDate(date):
        months = {1: 'janeiro', 2 : 'fevereiro', 3: 'março', 4: 'abril', 5: 'maio', 6: 'junho', 7: 'julho', 8: 'agosto', 9: 'setembro', 10: 'outubro', 11: 'novembro', 12: 'dezembro'}
        return '%s de %s de %s' % (date.day, months[int(date.month)], date.year)
